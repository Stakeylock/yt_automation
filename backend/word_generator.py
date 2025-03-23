from docx import Document

def generate_word_document(idea, thumbnail_desc, filming_suggestion, editing_flow, publishing_suggestions):
    doc = Document()
    doc.add_heading("Video Production Plan", level=1)
    
    doc.add_heading("1. Video Idea", level=2)
    doc.add_paragraph(idea)
    
    doc.add_heading("2. Thumbnail Description", level=2)
    doc.add_paragraph(thumbnail_desc)
    
    doc.add_heading("3. Filming Approach", level=2)
    doc.add_paragraph(filming_suggestion)
    
    doc.add_heading("4. Editing Flow", level=2)
    doc.add_paragraph(editing_flow)
    
    doc.add_heading("5. Publishing Suggestions", level=2)
    doc.add_paragraph(publishing_suggestions)
    
    file_path = "video_production_plan.docx"
    doc.save(file_path)
    return file_path
